"""
Convertly — File Converter  v2.0
Developer: Ibrahim Ezzeldin Mirghani
"""
import tkinter as tk
from tkinter import filedialog, ttk
import threading
import os
import subprocess
import platform
from datetime import datetime


def _open_path(path):
    if platform.system() == "Windows":
        os.startfile(path)
    elif platform.system() == "Darwin":
        subprocess.run(["open", path], check=False)
    else:
        subprocess.run(["xdg-open", path], check=False)

# ── Conversion Functions ─────────────────────────────────────────────────────

def _cleanup_docx_spacing(docx_path):
    """
    Post-process a converted DOCX to remove excessive blank space:
    - Collapses runs of more than 2 consecutive empty paragraphs into 1
    - Caps paragraph spaceBefore/spaceAfter values above 72 pt
    - Removes isolated page-break elements near the top of the document
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    try:
        doc = Document(docx_path)

        # ── 1. Remove excessive consecutive blank paragraphs ──
        blank_run = 0
        to_remove = []
        for para in doc.paragraphs:
            if not para.text.strip():
                blank_run += 1
                if blank_run > 2:
                    to_remove.append(para)
            else:
                blank_run = 0
        for para in to_remove:
            p = para._element
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)

        # ── 2. Cap runaway spaceBefore / spaceAfter ──
        for para in doc.paragraphs:
            pf = para.paragraph_format
            try:
                if pf.space_before and pf.space_before.pt > 72:
                    pf.space_before = Pt(18)
            except Exception:
                pass
            try:
                if pf.space_after and pf.space_after.pt > 72:
                    pf.space_after = Pt(8)
            except Exception:
                pass

        # ── 3. Remove stray page-break runs near the top ──
        for para in list(doc.paragraphs)[:10]:
            if para.text.strip():
                continue
            for run in para.runs:
                for br in run._element.findall(f'.//{qn("w:br")}'):
                    if br.get(qn('w:type')) == 'page':
                        br.getparent().remove(br)

        doc.save(docx_path)
    except Exception:
        pass  # never let cleanup crash the whole conversion


def pdf_to_word(pdf_path, out_path, stop_event=None):
    """
    Converts PDF → DOCX using the best available engine:
      1. Microsoft Word via COM (Windows, highest fidelity)
      2. LibreOffice via subprocess (if installed)
      3. pdf2docx (pure-Python fallback)
    Then post-processes to collapse excessive blank space.
    """
    import shutil
    abs_pdf = os.path.abspath(pdf_path)
    abs_out = os.path.abspath(out_path)
    converted = False

    # ── 1. Microsoft Word (COM) ───────────────────────────────────────────────
    if not converted:
        try:
            import pythoncom, win32com.client
            pythoncom.CoInitialize()
            word = None
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                doc = word.Documents.Open(abs_pdf)
                doc.SaveAs2(abs_out, FileFormat=16)
                doc.Close(False)
                converted = True
            finally:
                try:
                    if word: word.Quit()
                except Exception:
                    pass
                pythoncom.CoUninitialize()
        except Exception:
            pass

    # ── 2. LibreOffice ────────────────────────────────────────────────────────
    if not converted:
        lo_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for lo in lo_paths:
            if os.path.exists(lo):
                try:
                    out_dir = os.path.dirname(abs_out)
                    subprocess.run(
                        [lo, "--headless", "--convert-to", "docx",
                         "--outdir", out_dir, abs_pdf],
                        timeout=120, check=True,
                        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                    )
                    lo_out = os.path.join(
                        out_dir,
                        os.path.splitext(os.path.basename(pdf_path))[0] + ".docx"
                    )
                    if os.path.exists(lo_out) and os.path.abspath(lo_out) != abs_out:
                        shutil.move(lo_out, abs_out)
                    converted = True
                except Exception:
                    pass
                break

    # ── 3. pdf2docx (fallback) ────────────────────────────────────────────────
    if not converted:
        from pdf2docx import Converter
        cv = Converter(abs_pdf)
        cv.convert(abs_out)
        cv.close()

    # ── Post-process: collapse excessive blank space ───────────────────────────
    _cleanup_docx_spacing(abs_out)

def pdf_to_excel(pdf_path, out_path, stop_event=None):
    import pdfplumber, openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            if stop_event and stop_event.is_set():
                raise InterruptedError("Cancelled by user.")
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        ws.append([c if c else "" for c in row])
            else:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        ws.append([line])
    wb.save(out_path)

def word_to_pdf(docx_path, out_path, stop_event=None):
    """
    Converts DOCX to PDF using python-docx + reportlab.
    Preserves: document element order, inline images, text alignment,
    text colors, font sizes, bold/italic/underline, list bullets,
    and table cell background colors from the original DOCX.
    """
    import tempfile
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle,
                                    Image as RLImage)

    A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

    doc = Document(docx_path)
    pdf = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm,  bottomMargin=2.5*cm,
    )

    base_styles = getSampleStyleSheet()
    story      = []
    tmp_images = []
    max_w      = A4[0] - 5 * cm
    _sc        = [0]  # style name counter for uniqueness

    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        ALIGN_MAP = {
            WD_ALIGN_PARAGRAPH.LEFT:    TA_LEFT,
            WD_ALIGN_PARAGRAPH.CENTER:  TA_CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT:   TA_RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY,
        }
    except Exception:
        ALIGN_MAP = {}

    def esc(t):
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def run_markup(run):
        rt = esc(run.text)
        if not rt:
            return ""
        # Font color
        try:
            if run.font.color and run.font.color.type is not None:
                rgb = str(run.font.color.rgb)
                rt = f'<font color="#{rgb}">{rt}</font>'
        except Exception:
            pass
        # Font size
        try:
            if run.font.size:
                rt = f'<font size="{run.font.size.pt:.1f}">{rt}</font>'
        except Exception:
            pass
        # Bold / italic / underline
        if getattr(run, 'bold', False) and getattr(run, 'italic', False):
            rt = f"<b><i>{rt}</i></b>"
        elif getattr(run, 'bold', False):
            rt = f"<b>{rt}</b>"
        elif getattr(run, 'italic', False):
            rt = f"<i>{rt}</i>"
        if getattr(run, 'underline', False):
            rt = f"<u>{rt}</u>"
        return rt

    def make_style(style_name, alignment, base_size=11):
        _sc[0] += 1
        sn = f"_S{_sc[0]}"
        al = ALIGN_MAP.get(alignment, TA_LEFT)
        leading = max(base_size * 1.45, 14)
        if "Heading 1" in style_name:
            return ParagraphStyle(sn, parent=base_styles["Heading1"],
                fontSize=18, leading=22, spaceBefore=12, spaceAfter=6, alignment=al)
        elif "Heading 2" in style_name:
            return ParagraphStyle(sn, parent=base_styles["Heading2"],
                fontSize=14, leading=18, spaceBefore=10, spaceAfter=4, alignment=al)
        elif "Heading 3" in style_name:
            return ParagraphStyle(sn, parent=base_styles["Heading3"],
                fontSize=12, leading=16, spaceBefore=8, spaceAfter=3, alignment=al)
        else:
            return ParagraphStyle(sn, parent=base_styles["Normal"],
                fontSize=base_size, leading=leading, spaceAfter=4,
                leftIndent=(18 if "List" in style_name else 0),
                alignment=al)

    def extract_images(para_elem):
        imgs = []
        blips = para_elem.findall(f'.//{{{A_NS}}}blip')
        for blip in blips:
            r_embed = blip.get(f'{{{R_NS}}}embed')
            if not r_embed or r_embed not in doc.part.rels:
                continue
            try:
                img_part = doc.part.rels[r_embed].target_part
                img_data = img_part.blob
                ext = img_part.content_type.split('/')[-1]
                if ext == 'jpeg':
                    ext = 'jpg'
                tmp = tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False)
                tmp.write(img_data)
                tmp.close()
                tmp_images.append(tmp.name)
                # Get dimensions from EMU → pt
                extents = para_elem.findall(f'.//{{{WP_NS}}}extent')
                if extents:
                    cx = int(extents[0].get('cx', 0))
                    cy = int(extents[0].get('cy', 0))
                    w = cx / 914400 * 72
                    h = cy / 914400 * 72
                    if w > max_w:
                        h = h * max_w / w
                        w = max_w
                    img = RLImage(tmp.name, width=w, height=h)
                else:
                    img = RLImage(tmp.name, width=min(300, max_w))
                imgs.append(img)
            except Exception:
                pass
        return imgs

    def get_cell_bg(cell):
        tc_pr = cell._tc.find(qn('w:tcPr'))
        if tc_pr is not None:
            shd = tc_pr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill and fill not in ('auto',) and len(fill) == 6:
                    try:
                        return colors.HexColor(f"#{fill}")
                    except Exception:
                        pass
        return None

    # Iterate body elements in document order
    for elem in doc.element.body:
        if stop_event and stop_event.is_set():
            raise InterruptedError("Cancelled by user.")

        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

        if tag == 'p':
            para = DocxParagraph(elem, doc)

            # Embed any inline images first
            for img in extract_images(elem):
                story.append(img)
                story.append(Spacer(1, 4))

            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            style_name = para.style.name if para.style else "Normal"

            # Detect base font size from first sized run
            base_size = 11
            for run in para.runs:
                try:
                    if run.font.size:
                        base_size = run.font.size.pt
                        break
                except Exception:
                    pass

            p_style = make_style(style_name, para.alignment, base_size)

            parts = [run_markup(run) for run in para.runs]
            rich_text = "".join(parts) or esc(text)

            if "List Bullet" in style_name:
                rich_text = "• " + rich_text

            story.append(Paragraph(rich_text, p_style))

        elif tag == 'tbl':
            tbl = DocxTable(elem, doc)
            data   = []
            bg_map = {}  # (row_idx, col_idx) → HexColor

            for r_idx, row in enumerate(tbl.rows):
                row_data = []
                for c_idx, cell in enumerate(row.cells):
                    row_data.append(cell.text)
                    bg = get_cell_bg(cell)
                    if bg:
                        bg_map[(r_idx, c_idx)] = bg
                data.append(row_data)

            if not data:
                continue

            num_cols = max(len(r) for r in data)
            col_w    = max_w / max(num_cols, 1)
            data     = [r + [""] * (num_cols - len(r)) for r in data]

            ts = [
                ("FONTSIZE",       (0, 0), (-1, -1), 9),
                ("GRID",           (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                ("TOPPADDING",     (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
                ("LEFTPADDING",    (0, 0), (-1, -1), 6),
                ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
            ]
            for (r, c), bg in bg_map.items():
                ts.append(("BACKGROUND", (c, r), (c, r), bg))

            # Default header row only when the first cell has no custom color
            if (0, 0) not in bg_map:
                ts += [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4361EE")),
                    ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                    ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                     [colors.white, colors.HexColor("#F5F7FF")]),
                ]

            rl_tbl = Table(data, colWidths=[col_w] * num_cols)
            rl_tbl.setStyle(TableStyle(ts))
            story.append(Spacer(1, 6))
            story.append(rl_tbl)
            story.append(Spacer(1, 6))

    if not story:
        story.append(Paragraph("(Empty document)",
            ParagraphStyle("_empty", parent=base_styles["Normal"], fontSize=11)))

    try:
        pdf.build(story)
    finally:
        for tmp in tmp_images:
            try:
                os.remove(tmp)
            except Exception:
                pass

def excel_to_pdf(xlsx_path, out_path, stop_event=None):
    """
    Uses openpyxl + reportlab to convert Excel → PDF.
    Preserves: cell background colors, font bold, cell alignment,
    and merged cell spans from the original spreadsheet.
    """
    import openpyxl
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    wb    = openpyxl.load_workbook(xlsx_path, data_only=True)
    story = []
    styles = getSampleStyleSheet()

    for sheet_name in wb.sheetnames:
        if stop_event and stop_event.is_set():
            raise InterruptedError("Cancelled by user.")

        ws = wb[sheet_name]

        story.append(Paragraph(f"<b>{sheet_name}</b>", styles["Heading2"]))
        story.append(Spacer(1, 0.3 * cm))

        rows = list(ws.iter_rows())
        if not rows:
            continue

        num_cols = ws.max_column or 1
        num_rows = ws.max_row    or 1

        # Base table commands
        ts = [
            ("FONTSIZE",       (0, 0), (-1, -1), 8),
            ("ALIGN",          (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",     (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 3),
            ("LEFTPADDING",    (0, 0), (-1, -1), 4),
            ("GRID",           (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
        ]

        data        = []
        has_any_bg  = False

        for r_idx, row in enumerate(rows):
            row_data = []
            for c_idx, cell in enumerate(row):
                val = str(cell.value) if cell.value is not None else ""
                row_data.append(val)

                # Cell background color
                try:
                    fill = cell.fill
                    if fill and fill.fill_type not in (None, 'none'):
                        fg = fill.fgColor
                        if fg and fg.type == 'rgb' and fg.rgb:
                            rgb = fg.rgb[-6:]  # strip alpha channel
                            if rgb.upper() not in ('FFFFFF', '000000', '000000'):
                                ts.append(("BACKGROUND",
                                           (c_idx, r_idx), (c_idx, r_idx),
                                           colors.HexColor(f"#{rgb}")))
                                has_any_bg = True
                except Exception:
                    pass

                # Font bold
                try:
                    if cell.font and cell.font.bold:
                        ts.append(("FONTNAME",
                                   (c_idx, r_idx), (c_idx, r_idx),
                                   "Helvetica-Bold"))
                except Exception:
                    pass

                # Cell text alignment
                try:
                    if cell.alignment and cell.alignment.horizontal:
                        al_map = {
                            'center':  'CENTER',
                            'right':   'RIGHT',
                            'left':    'LEFT',
                            'general': 'LEFT',
                        }
                        al = al_map.get(cell.alignment.horizontal)
                        if al:
                            ts.append(("ALIGN",
                                       (c_idx, r_idx), (c_idx, r_idx), al))
                except Exception:
                    pass

            data.append(row_data)

        if not any(any(c for c in r) for r in data):
            continue

        # Merged cell spans
        for merge in ws.merged_cells.ranges:
            r1 = merge.min_row - 1
            c1 = merge.min_col - 1
            r2 = merge.max_row - 1
            c2 = merge.max_col - 1
            ts.append(("SPAN", (c1, r1), (c2, r2)))

        # Fall back to default header styling only when no cell has a custom color
        if not has_any_bg:
            ts += [
                ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#4361EE")),
                ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
                ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F5F7FF")]),
            ]

        # Pad rows and build table
        data = [r + [""] * (num_cols - len(r)) for r in data]

        page_w  = landscape(A4)[0] - 2 * cm
        col_w   = page_w / max(num_cols, 1)

        tbl = Table(data, colWidths=[col_w] * num_cols, repeatRows=1)
        tbl.setStyle(TableStyle(ts))
        story.append(tbl)
        story.append(Spacer(1, 0.5 * cm))

    doc = SimpleDocTemplate(
        out_path,
        pagesize=landscape(A4),
        leftMargin=1*cm, rightMargin=1*cm,
        topMargin=1*cm,  bottomMargin=1*cm,
    )
    doc.build(story)


# ── Design Tokens ─────────────────────────────────────────────────────────────
BG          = "#F8F9FB"
WHITE       = "#FFFFFF"
BORDER      = "#EAEDF3"
BORDER_SOFT = "#F0F2F7"
TEXT        = "#0F1117"
TEXT_SEC    = "#6B7280"
TEXT_MUTED  = "#9CA3AF"
ACCENT      = "#4361EE"
ACCENT_DARK = "#3451D1"
ACCENT_SOFT = "#EEF1FF"

MODES = [
    {
        "label": "PDF → Word",
        "desc":  "Editable Word document",
        "icon":  "📄",
        "color": "#4361EE",
        "light": "#EEF1FF",
        "ft":    [("PDF Files", "*.pdf")],
        "ext":   "_converted.docx",
        "fn":    pdf_to_word,
    },
    {
        "label": "PDF → Excel",
        "desc":  "Extract tables & data",
        "icon":  "📊",
        "color": "#10B981",
        "light": "#ECFDF5",
        "ft":    [("PDF Files", "*.pdf")],
        "ext":   "_converted.xlsx",
        "fn":    pdf_to_excel,
    },
    {
        "label": "Word → PDF",
        "desc":  "Professional PDF output",
        "icon":  "📝",
        "color": "#EF4444",
        "light": "#FFF1F2",
        "ft":    [("Word Files", "*.docx")],
        "ext":   "_converted.pdf",
        "fn":    word_to_pdf,
    },
    {
        "label": "Excel → PDF",
        "desc":  "Spreadsheet to PDF",
        "icon":  "📈",
        "color": "#F59E0B",
        "light": "#FFFBEB",
        "ft":    [("Excel Files", "*.xlsx")],
        "ext":   "_converted.pdf",
        "fn":    excel_to_pdf,
    },
]


class ConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Convertly — File Converter")
        self.root.geometry("620x900")
        self.root.minsize(540, 820)
        self.root.resizable(True, True)
        self.root.configure(bg=BG)

        # Center on screen
        self.root.update_idletasks()
        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()
        self.root.geometry(f"620x900+{(sw-620)//2}+{(sh-900)//2}")

        self.file_path   = tk.StringVar()
        self.active_mode = MODES[0]
        self.card_refs   = []
        self.status_var  = tk.StringVar(value="Choose a format, then select your file.")
        self._stop_event = threading.Event()
        self._converting = False

        self._build()

    # ─────────────────────────────────────────────────────────────────────────
    def _build(self):
        self._build_header()
        self._build_body()
        self._build_footer()
        self._activate_card(0)

    # ── Header ───────────────────────────────────────────────────────────────
    def _build_header(self):
        hdr = tk.Frame(self.root, bg=WHITE)
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=BORDER, height=1).pack(side="bottom", fill="x")

        inner = tk.Frame(hdr, bg=WHITE)
        inner.pack(fill="x", padx=32, pady=18)

        logo = tk.Frame(inner, bg=WHITE)
        logo.pack(side="left")

        badge = tk.Frame(logo, bg=ACCENT, width=40, height=40)
        badge.pack(side="left")
        badge.pack_propagate(False)
        tk.Label(badge, text="CV", font=("Segoe UI", 12, "bold"),
                 bg=ACCENT, fg=WHITE).pack(expand=True)

        name_col = tk.Frame(logo, bg=WHITE)
        name_col.pack(side="left", padx=(12, 0))
        tk.Label(name_col, text="Convertly",
                 font=("Segoe UI", 15, "bold"),
                 bg=WHITE, fg=TEXT).pack(anchor="w")
        tk.Label(name_col, text="File Format Converter",
                 font=("Segoe UI", 8),
                 bg=WHITE, fg=TEXT_MUTED).pack(anchor="w")

        pill = tk.Frame(inner, bg=ACCENT_SOFT, padx=12, pady=5)
        pill.pack(side="right", anchor="center")
        tk.Label(pill, text="v 2.1", font=("Segoe UI", 8, "bold"),
                 bg=ACCENT_SOFT, fg=ACCENT).pack()

    # ── Body ─────────────────────────────────────────────────────────────────
    def _build_body(self):
        body = tk.Frame(self.root, bg=BG)
        body.pack(fill="both", expand=True, padx=28, pady=16)
        body.columnconfigure(0, weight=1)

        self._section(body, "Conversion Type")

        grid = tk.Frame(body, bg=BG)
        grid.pack(fill="x")
        grid.columnconfigure(0, weight=1)
        grid.columnconfigure(1, weight=1)

        for i, mode in enumerate(MODES):
            self._make_card(grid, mode, i // 2, i % 2)

        self._section(body, "Select File", top=16)

        fzone = tk.Frame(body, bg=WHITE,
                         highlightthickness=1,
                         highlightbackground=BORDER)
        fzone.pack(fill="x")

        fz_in = tk.Frame(fzone, bg=WHITE)
        fz_in.pack(fill="x", padx=18, pady=14)

        self.file_icon_lbl = tk.Label(fz_in, text="📂",
                                      font=("Segoe UI", 22), bg=WHITE)
        self.file_icon_lbl.pack(side="left", padx=(0, 14))

        txt = tk.Frame(fz_in, bg=WHITE)
        txt.pack(side="left", fill="both", expand=True)

        self.fname_lbl = tk.Label(txt, text="No file selected",
                                  font=("Segoe UI", 10, "bold"),
                                  bg=WHITE, fg=TEXT_SEC, anchor="w")
        self.fname_lbl.pack(fill="x")

        self.fpath_lbl = tk.Label(txt, text="Supports PDF · DOCX · XLSX",
                                  font=("Segoe UI", 8),
                                  bg=WHITE, fg=TEXT_MUTED, anchor="w")
        self.fpath_lbl.pack(fill="x", pady=(2, 0))

        self.browse_btn = tk.Button(fz_in, text="Browse",
                                    command=self.browse,
                                    bg=ACCENT, fg=WHITE,
                                    font=("Segoe UI", 9, "bold"),
                                    relief="flat", padx=18, pady=8,
                                    cursor="hand2", bd=0,
                                    activebackground=ACCENT_DARK,
                                    activeforeground=WHITE)
        self.browse_btn.pack(side="right")
        self._btn_hover(self.browse_btn, ACCENT, ACCENT_DARK)

        # ── Convert + Stop buttons row ──
        btn_row = tk.Frame(body, bg=BG)
        btn_row.pack(fill="x", pady=(20, 0))
        btn_row.columnconfigure(0, weight=1)

        self.convert_btn = tk.Button(btn_row, text="  Convert Now  →",
                                     command=self.start_conversion,
                                     bg=ACCENT, fg=WHITE,
                                     font=("Segoe UI", 12, "bold"),
                                     relief="flat", pady=14,
                                     cursor="hand2", bd=0,
                                     activebackground=ACCENT_DARK,
                                     activeforeground=WHITE)
        self.convert_btn.pack(side="left", fill="x", expand=True)
        self._btn_hover(self.convert_btn, ACCENT, ACCENT_DARK)

        # Stop button — hidden until conversion starts
        self.stop_btn = tk.Button(btn_row, text="✕  Stop",
                                  command=self.stop_conversion,
                                  bg="#EF4444", fg=WHITE,
                                  font=("Segoe UI", 11, "bold"),
                                  relief="flat", pady=14, padx=20,
                                  cursor="hand2", bd=0,
                                  activebackground="#C0392B",
                                  activeforeground=WHITE)
        self._btn_hover(self.stop_btn, "#EF4444", "#C0392B")

        # Progress
        style = ttk.Style()
        style.theme_use("default")
        style.configure("App.Horizontal.TProgressbar",
                        troughcolor=BORDER_SOFT,
                        background=ACCENT,
                        thickness=3)

        self.progress = ttk.Progressbar(body, mode="indeterminate",
                                        style="App.Horizontal.TProgressbar")
        self.progress.pack(fill="x", pady=(10, 0))

        tk.Label(body, textvariable=self.status_var,
                 font=("Segoe UI", 9),
                 bg=BG, fg=TEXT_MUTED, anchor="w").pack(fill="x", pady=(6, 0))

    # ── Footer ───────────────────────────────────────────────────────────────
    def _build_footer(self):
        year = datetime.now().year

        footer = tk.Frame(self.root, bg=WHITE)
        footer.pack(fill="x", side="bottom")
        tk.Frame(footer, bg=BORDER, height=1).pack(fill="x")

        f_in = tk.Frame(footer, bg=WHITE)
        f_in.pack(fill="x", padx=28, pady=10)

        tk.Label(f_in,
                 text="🔒  Converted locally — your files stay private.",
                 font=("Segoe UI", 8), bg=WHITE, fg=TEXT_MUTED,
                 anchor="w").pack(side="left")

        dev = tk.Frame(f_in, bg=ACCENT_SOFT,
                       highlightthickness=1, highlightbackground="#D8DEFF")
        dev.pack(side="right")
        dev_in = tk.Frame(dev, bg=ACCENT_SOFT)
        dev_in.pack(padx=12, pady=6)
        tk.Label(dev_in, text="Ibrahim Ezzeldin Mirghani",
                 font=("Segoe UI", 8, "bold"),
                 bg=ACCENT_SOFT, fg=TEXT).pack(anchor="e")
        tk.Label(dev_in, text="Application's Developer",
                 font=("Consolas", 7),
                 bg=ACCENT_SOFT, fg=ACCENT).pack(anchor="e")

        copy_bar = tk.Frame(self.root, bg=BORDER_SOFT)
        copy_bar.pack(fill="x", side="bottom")
        tk.Label(copy_bar,
                 text=f"© {year}  Convertly — File Converter  •  All rights reserved.",
                 font=("Segoe UI", 8),
                 bg=BORDER_SOFT, fg=TEXT_MUTED).pack(pady=5)

    # ── Helpers ──────────────────────────────────────────────────────────────
    def _section(self, parent, text, top=0):
        tk.Label(parent, text=text,
                 font=("Segoe UI", 10, "bold"),
                 bg=BG, fg=TEXT, anchor="w").pack(fill="x", pady=(top, 8))

    def _btn_hover(self, btn, normal, hover):
        btn.bind("<Enter>", lambda e: btn.configure(bg=hover))
        btn.bind("<Leave>", lambda e: btn.configure(bg=normal))

    def _make_card(self, parent, mode, row, col):
        pad = (0, 6) if col == 0 else (6, 0)

        outer = tk.Frame(parent, bg=BORDER, cursor="hand2")
        outer.grid(row=row, column=col, sticky="ew", padx=pad, pady=(0, 10))

        card = tk.Frame(outer, bg=WHITE, cursor="hand2", height=110)
        card.pack(fill="both", padx=1, pady=1)
        card.pack_propagate(False)

        top_row = tk.Frame(card, bg=WHITE)
        top_row.pack(fill="x", padx=14, pady=(13, 4))

        icon_frame = tk.Frame(top_row, bg=mode["light"], width=34, height=34)
        icon_frame.pack(side="left")
        icon_frame.pack_propagate(False)
        tk.Label(icon_frame, text=mode["icon"],
                 font=("Segoe UI", 15), bg=mode["light"]).pack(expand=True)

        check = tk.Label(top_row, text="",
                         font=("Segoe UI", 13),
                         bg=WHITE, fg=mode["color"])
        check.pack(side="right")

        dot = tk.Frame(top_row, bg=mode["color"], width=7, height=7)
        dot.pack(side="right", padx=(0, 4))

        tk.Label(card, text=mode["label"],
                 font=("Segoe UI", 10, "bold"),
                 bg=WHITE, fg=TEXT, anchor="w").pack(fill="x", padx=14)

        tk.Label(card, text=mode["desc"],
                 font=("Segoe UI", 8),
                 bg=WHITE, fg=TEXT_MUTED, anchor="w").pack(fill="x", padx=14)

        idx = len(self.card_refs)
        self.card_refs.append({
            "outer": outer, "card": card,
            "check": check, "top_row": top_row,
            "icon_frame": icon_frame, "dot": dot,
            "mode": mode
        })

        all_widgets = [outer, card, top_row, icon_frame, check, dot] + \
                      list(card.winfo_children()) + \
                      list(top_row.winfo_children()) + \
                      list(icon_frame.winfo_children())

        for w in all_widgets:
            w.bind("<Button-1>", lambda e, i=idx: self._activate_card(i))
            w.bind("<Enter>",    lambda e, i=idx: self._card_hover(i, True))
            w.bind("<Leave>",    lambda e, i=idx: self._card_hover(i, False))

    def _card_hover(self, idx, entering):
        ref = self.card_refs[idx]
        if self.active_mode != ref["mode"]:
            ref["outer"].configure(bg=BORDER_SOFT if entering else BORDER)

    def _activate_card(self, idx):
        if self._converting:
            return  # don't allow switching while converting

        for ref in self.card_refs:
            ref["outer"].configure(bg=BORDER)
            ref["card"].configure(bg=WHITE)
            ref["check"].configure(text="", bg=WHITE)
            ref["top_row"].configure(bg=WHITE)
            ref["icon_frame"].configure(bg=ref["mode"]["light"])
            for child in list(ref["card"].winfo_children()) + \
                         list(ref["top_row"].winfo_children()) + \
                         list(ref["icon_frame"].winfo_children()):
                try:    child.configure(bg=WHITE)
                except: pass

        ref  = self.card_refs[idx]
        mode = ref["mode"]
        ref["outer"].configure(bg=mode["color"])
        ref["card"].configure(bg=mode["light"])
        ref["check"].configure(text="✓", bg=mode["light"])
        ref["top_row"].configure(bg=mode["light"])
        ref["icon_frame"].configure(bg=mode["light"])
        for child in list(ref["card"].winfo_children()) + \
                     list(ref["top_row"].winfo_children()) + \
                     list(ref["icon_frame"].winfo_children()):
            try:    child.configure(bg=mode["light"])
            except: pass

        self.active_mode = mode

        for btn in [self.convert_btn, self.browse_btn]:
            if btn:
                btn.configure(bg=mode["color"], activebackground=mode["color"])
                self._btn_hover(btn, mode["color"], self._darken(mode["color"]))

        style = ttk.Style()
        style.configure("App.Horizontal.TProgressbar",
                        background=mode["color"])

    def _darken(self, hex_color, factor=0.85):
        hex_color = hex_color.lstrip("#")
        r, g, b = (int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        return "#{:02x}{:02x}{:02x}".format(
            int(r * factor), int(g * factor), int(b * factor))

    # ── File & Conversion ────────────────────────────────────────────────────
    def browse(self):
        path = filedialog.askopenfilename(filetypes=self.active_mode["ft"])
        if path:
            self.file_path.set(path)
            name = os.path.basename(path)
            self.fname_lbl.config(
                text=name if len(name) <= 42 else name[:39] + "…",
                fg=TEXT)
            self.fpath_lbl.config(
                text=path if len(path) <= 58 else "…" + path[-55:],
                fg=TEXT_MUTED)
            self.file_icon_lbl.config(text="📄")
            self.status_var.set("File ready — click Convert Now to proceed.")

    def start_conversion(self):
        if not self.file_path.get():
            self._error_popup("No file selected. Please click Browse first.")
            return

        self._stop_event.clear()
        self._converting = True

        # Show stop button
        self.convert_btn.config(state="disabled", text="  Converting…  ⏳")
        self.stop_btn.pack(side="left", padx=(8, 0))

        threading.Thread(target=self._do_convert, daemon=True).start()

    def stop_conversion(self):
        self._stop_event.set()
        self.status_var.set("⚠  Stopping — please wait…")
        self.stop_btn.config(state="disabled", text="Stopping…")

    def _do_convert(self):
        self.progress.start(8)
        self.status_var.set("Converting — please wait…")
        try:
            src  = self.file_path.get()
            base = os.path.splitext(src)[0]
            out  = base + self.active_mode["ext"]
            self.active_mode["fn"](src, out, self._stop_event)

            if self._stop_event.is_set():
                # Clean up partial output
                if os.path.exists(out):
                    try: os.remove(out)
                    except: pass
                self._finish_ui("⚠  Conversion stopped.")
            else:
                self._finish_ui(f"✓  Done!  Saved as {os.path.basename(out)}", out)

        except InterruptedError:
            self._finish_ui("⚠  Conversion cancelled.")
        except Exception as e:
            self._finish_ui("An error occurred during conversion.", error=str(e))

    def _finish_ui(self, status_msg, out_path=None, error=None):
        self.progress.stop()
        self._converting = False
        self._stop_event.clear()

        self.status_var.set(status_msg)
        self.convert_btn.config(state="normal", text="  Convert Now  →")
        self.stop_btn.config(state="normal", text="✕  Stop")

        # Hide stop button
        self.stop_btn.pack_forget()

        if out_path:
            captured = out_path
            self.root.after(0, lambda p=captured: self._success_popup(p))
        elif error:
            captured = error
            self.root.after(0, lambda e=captured: self._error_popup(e))

    # ── Popups ───────────────────────────────────────────────────────────────
    def _center_popup(self, popup, w, h):
        popup.update_idletasks()
        x = self.root.winfo_x() + self.root.winfo_width()  // 2 - w // 2
        y = self.root.winfo_y() + self.root.winfo_height() // 2 - h // 2
        popup.geometry(f"{w}x{h}+{x}+{y}")

    def _success_popup(self, out_path):
        p = tk.Toplevel(self.root)
        p.title("Done!")
        p.resizable(False, False)
        p.configure(bg=WHITE)
        p.grab_set()
        self._center_popup(p, 480, 340)

        tk.Frame(p, bg=self.active_mode["color"], height=4).pack(fill="x")

        body = tk.Frame(p, bg=WHITE)
        body.pack(fill="both", expand=True, padx=30, pady=20)

        # Top: badge + title side by side
        top = tk.Frame(body, bg=WHITE)
        top.pack(fill="x", pady=(0, 12))

        badge = tk.Frame(top, bg=self.active_mode["light"], width=46, height=46)
        badge.pack(side="left")
        badge.pack_propagate(False)
        tk.Label(badge, text="✓",
                 font=("Segoe UI", 17, "bold"),
                 bg=self.active_mode["light"],
                 fg=self.active_mode["color"]).pack(expand=True)

        title_col = tk.Frame(top, bg=WHITE)
        title_col.pack(side="left", padx=(14, 0))
        tk.Label(title_col, text="Conversion complete!",
                 font=("Segoe UI", 13, "bold"),
                 bg=WHITE, fg=TEXT, anchor="w").pack(anchor="w")
        tk.Label(title_col, text="Your file has been saved successfully.",
                 font=("Segoe UI", 8),
                 bg=WHITE, fg=TEXT_MUTED, anchor="w").pack(anchor="w")

        # File info box
        info_box = tk.Frame(body, bg=BORDER_SOFT,
                            highlightthickness=1,
                            highlightbackground=BORDER)
        info_box.pack(fill="x", pady=(0, 16))
        info_in = tk.Frame(info_box, bg=BORDER_SOFT)
        info_in.pack(fill="x", padx=12, pady=10)

        tk.Label(info_in, text="📄  " + os.path.basename(out_path),
                 font=("Segoe UI", 10, "bold"),
                 bg=BORDER_SOFT, fg=TEXT, anchor="w").pack(fill="x")
        tk.Label(info_in, text=out_path,
                 font=("Segoe UI", 8),
                 bg=BORDER_SOFT, fg=TEXT_MUTED, anchor="w",
                 wraplength=400).pack(fill="x", pady=(3, 0))

        tk.Frame(body, bg=BORDER, height=1).pack(fill="x")

        # Buttons — always visible, full width
        btns = tk.Frame(body, bg=WHITE)
        btns.pack(fill="x", pady=(14, 0))
        btns.columnconfigure(0, weight=1)
        btns.columnconfigure(1, weight=1)
        btns.columnconfigure(2, weight=1)

        c = self.active_mode["color"]
        tk.Button(btns, text="📂  Open File",
                  command=lambda: (_open_path(out_path), p.destroy()),
                  bg=c, fg=WHITE, font=("Segoe UI", 10, "bold"),
                  relief="flat", pady=10,
                  cursor="hand2", bd=0).grid(row=0, column=0, sticky="ew", padx=(0, 4))

        tk.Button(btns, text="📁  Open Folder",
                  command=lambda: (_open_path(os.path.dirname(out_path)), p.destroy()),
                  bg=BORDER_SOFT, fg=TEXT, font=("Segoe UI", 10),
                  relief="flat", pady=10,
                  cursor="hand2", bd=0).grid(row=0, column=1, sticky="ew", padx=4)

        tk.Button(btns, text="✕  Close", command=p.destroy,
                  bg=WHITE, fg=TEXT_MUTED, font=("Segoe UI", 10),
                  relief="flat", pady=10,
                  cursor="hand2", bd=0,
                  highlightthickness=1,
                  highlightbackground=BORDER).grid(row=0, column=2, sticky="ew", padx=(4, 0))

    def _error_popup(self, message):
        p = tk.Toplevel(self.root)
        p.title("Error")
        p.resizable(False, False)
        p.configure(bg=WHITE)
        p.grab_set()
        self._center_popup(p, 420, 220)

        tk.Frame(p, bg="#EF4444", height=4).pack(fill="x")

        body = tk.Frame(p, bg=WHITE)
        body.pack(fill="both", expand=True, padx=30, pady=24)

        tk.Label(body, text="Something went wrong",
                 font=("Segoe UI", 13, "bold"),
                 bg=WHITE, fg=TEXT, anchor="w").pack(fill="x")

        tk.Label(body, text=message,
                 font=("Segoe UI", 9),
                 bg=WHITE, fg=TEXT_SEC, anchor="w",
                 wraplength=360, justify="left").pack(fill="x", pady=(8, 20))

        tk.Button(body, text="Close", command=p.destroy,
                  bg="#EF4444", fg=WHITE,
                  font=("Segoe UI", 10, "bold"),
                  relief="flat", padx=20, pady=8,
                  cursor="hand2", bd=0).pack(anchor="w")


# ── Run ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    root = tk.Tk()
    app  = ConverterApp(root)
    root.mainloop()